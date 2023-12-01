import re
import zipfile
import time
import pdfminer.high_level
import ebooklib
import tempfile
import os
import openai
import json
import random
import docx
import pandas as pd
import mobi
import sys

from docx import Document
from lxml import etree
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from ebooklib import epub
from ebooklib.epub import EpubImage
from bs4 import BeautifulSoup

class ProcessFile:
    """A class about according to the file extension, use the corresponding function to convert the file to text."""

    def __init__(self, parameterReader):
        """Initialize the title of filename and text which receives the contents of file."""
        self.filename = ""
        self.start_page = 0
        self.end_page = 0
        self.total_pages = 0
        self.transliteration_list_file = ""
        self.transliteration_word_capi_low = ""
        self.bilingual_output = ""
        self.prompt = ""
        self.language_code = ""
        self.jsonfile = ""
        self.translated_dict = ""
        self.new_filename = ""
        self.new_filenametxt = ""
        self.show = ""
        self.azure = ""
        self.tlist = ""
        self.test = ""
        self.gpt_model = ""
        self.gpt_temperature = ""

        self.title = ""
        self.text = ""
        self.book = ""
        self.total_tokens = 0
        self.completion_tokens = 0
        self.prompt_tokens = 0
        self.short_text_list = ""
        self.translated_text = ""
        self.translated_short_text = ""
        self.count = 0
        self.messages = ""

        self.client = ""
        self.non_azure_client = ""

        self._set_args_from_parameterReader(parameterReader)

    def _set_args_from_parameterReader(self, parameterReader):
        """Set args from parameterReader."""
        self.filename = parameterReader.filename
        self.start_page = parameterReader.startpage
        self.end_page = parameterReader.endpage
        self.total_pages = 0
        self.transliteration_list_file = parameterReader.transliteration_list_file
        self.transliteration_word_capi_low = parameterReader.transliteration_word_capi_low
        self.bilingual_output = parameterReader.bilingual_output
        self.prompt = parameterReader.prompt
        self.language_code = parameterReader.language_code
        self.jsonfile = parameterReader.jsonfile
        self.translated_dict = parameterReader.translated_dict
        self.new_filename = parameterReader.new_filename
        self.new_filenametxt = parameterReader.new_filenametxt
        self.show = parameterReader.show
        self.tlist = parameterReader.tlist
        self.test = parameterReader.test
        self.gpt_model = parameterReader.gpt_model
        self.gpt_temperature = parameterReader.gpt_temperature
        self.api_proxy = parameterReader.api_proxy

        self.azure = parameterReader.azure
        if self.azure:
            self.client = parameterReader.client
            self.openai_api_model_azure = parameterReader.openai_api_model_azure
        
        if len(self.api_proxy) != 0:
            self.non_azure_client = parameterReader.non_azure_client


    def _get_pdf_total_pages(self):
        """Get total pages."""
        with open(self.filename, 'rb') as file:
            parser = PDFParser(file)
            document = PDFDocument(parser)
            self.total_pages = len(list(PDFPage.create_pages(document)))

    def _convert_pdf_to_text(self):
        """Access the contents of the PDF file and convert it to text."""
        print("\033[1;32mINFO:Converting pdf to text.\033[0m")
        self.text = pdfminer.high_level.extract_text(self.filename, page_numbers=list(range(self.start_page - 1, self.end_page)))
        print("-" * 3)
        if self.show:
            print("*" * 3)
            print(self.text)
            print("*" * 3)
        print("\033[1;32mINFO:Finished converting pdf to text\033[0m")

    def _convert_mobi_to_text(self):
        """Access the content fo mobi and then convert it to text."""
        # Extract MOBI contents to a temporary directory
        with tempfile.TemporaryDirectory() as tempdir:
            tempdir, filepath = mobi.extract(self.filename)

            # Find the HTML file in the temporary directory
            for root, _, files in os.walk(tempdir):
                for file in files:
                    if file.endswith(".html"):
                        html_file = os.path.join(root, file)
                        break
                else:
                    continue
                break
            else:
                raise FileNotFoundError("ERROR:HTML file not found in the extracted MOBI contents")

            # Parse the HTML file with BeautifulSoup to get the text
            with open(html_file, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                self.text = soup.get_text()

    def _convert_docx_to_text(self):
        """Access the content of docx and then convert it to text."""
        print("-" * 3)
        print("\033[1;32mINFO:Parsing the DOCX content.\033[0m")
        doc = docx.Document(self.filename)

        for paragraph in doc.paragraphs:
            self.text += paragraph.text + "\n"

    def _convert_epub_to_text(self):
        """Convert epub to text."""
        # Access all contents
        for item in self.book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Use BeautifulSoup to extract the original text
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                self.text += re.sub(r'\n+', '\n', soup.get_text().strip())

    def _text_replace(self):
        """Replace the text according to the transliteration table."""
        # Read the excel file and store the first column and the second column as two lists
        df = pd.read_excel(self.transliteration_list_file)
        old_words = df.iloc[:, 0].tolist()
        new_words = df.iloc[:, 1].tolist()
        # Order the old word list in descending order of length and synchronize the new word list
        old_words, new_words = zip(*sorted(zip(old_words, new_words), key=lambda x: len(x[0]), reverse=True))
        # Iterate through two lists and replace strings
        for i in range(len(old_words)):
            # If ingore the case, convert the string and the word to be replaced to lowercase
            if not self.transliteration_word_capi_low:
                lower_old_word = old_words[i].lower()
                # Use the regular expression to replace, note that the original string case is retained
                self.text = re.sub(r"\b" + lower_old_word + r"\b", new_words[i], self.text, flags=re.IGNORECASE)
            else:
                # If care about the case, just use the regular expression to replace
                self.text = re.sub(r"\b" + old_words[i] + r"\b", new_words[i], self.text)

    def _text_replace_reverse(self, text):
        """Replace the text according to the transliteration table in reverse order."""
        # Read the excel file and store the first column and the second column as two lists
        df = pd.read_excel(self.transliteration_list_file)
        old_words = df.iloc[:, 0].tolist()  # Swapped
        new_words = df.iloc[:, 1].tolist()  # Swapped
        # Order the new word list in descending order of length and synchronize the old word list
        new_words, old_words = zip(*sorted(zip(new_words, old_words), key=lambda x: len(x[0]), reverse=True))
        # Iterate through two lists and replace strings
        for i in range(len(new_words)):
            # If ignore the case, convert the string and the word to be replaced to lowercase
            if not self.transliteration_word_capi_low:
                lower_new_word = new_words[i].lower()
                # Use the regular expression to replace, note that the original string case is retained
                text = re.sub(r"\b" + lower_new_word + r"\b", old_words[i], text, flags=re.IGNORECASE)
            else:
                # If care about the case, just use the regular expression to replace
                text = re.sub(r"\b" + new_words[i] + r"\b", old_words[i], text)

        return text

    def _reverse_text_replace_reverse(self, text):
        """Reverse the text according to the transliteration table in reverse order."""
        # Read the excel file and store the first column and the second column as two lists
        df = pd.read_excel(self.transliteration_list_file)
        new_words = df.iloc[:, 0].tolist()  # Swapped
        old_words = df.iloc[:, 1].tolist()  # Swapped
        # Order the new word list in descending order of length and synchronize the old word list
        new_words, old_words = zip(*sorted(zip(new_words, old_words), key=lambda x: len(x[0]), reverse=True))
        # Iterate through two lists and replace strings
        for i in range(len(new_words)):
            # If ignore the case, convert the string and the word to be replaced to lowercase
            if not self.transliteration_word_capi_low:
                lower_new_word = new_words[i].lower()
                # Use the regular expression to replace, note that the original string case is retained
                text = re.sub(r"\b" + lower_new_word + r"\b", old_words[i], text, flags=re.IGNORECASE)
            else:
                # If care about the case, just use the regular expression to replace
                text = re.sub(r"\b" + new_words[i] + r"\b", old_words[i], text)

        return text

    def _split_text(self):
        """Divide the text into a list of short texts with no more than 1024 characters."""
        # Use the regular expression to split the text into a list of sentences
        sentence_list = re.findall(r'.+?[。！？!?.]', self.text)
        # Initialize the short text list
        self.short_text_list = []
        # Initialize the current short text
        short_text = ""
        # Iterate through the sentence list
        for s in sentence_list:
            # If the current short plus the length of the new sentence is not greater than 1024, add the new sentence to the current short
            if len(short_text + s) <= 1024:
                short_text += s
            # If the current short plus the length of the new sentence is greater than 1024, add the current short to the short text list and reset the current short to the new sentence
            else:
                self.short_text_list.append(short_text)
                short_text = s
        # Add the last short text to the short text list
        self.short_text_list.append(short_text)

    def _replace_sign(self, text):
        """Replace the period with a period plus line break."""
        text = text.replace(". ", ".\n")
        text = text.replace("。", "。\n")
        text = text.replace("?", "?\n")
        text = text.replace("？", "？\n")
        text = text.replace("！", "！\n")
        text = text.replace("。\n”", "。”\n")
        text = text.replace("！\n”", "！”\n")
        text = text.replace("？\n”", "？”\n")
        return text

    def _get_completion_from_messages(self):
        """Get completion from messages."""
        if len(self.api_proxy) == 0:
            response = openai.chat.completions.create(
                model=self.gpt_model,
                messages=self.messages,
                temperature=self.gpt_temperature,
            )
        else:
            response = self.non_azure_client.chat.completions.create(
                model=self.gpt_model,
                messages=self.messages,
                temperature=self.gpt_temperature,
            )

        content = response.choices[0].message.content

        token_dict = {
        'prompt_tokens':response.usage.prompt_tokens,
        'completion_tokens':response.usage.completion_tokens,
        'total_tokens':response.usage.total_tokens,
        }

        return content, token_dict

    def _get_completion_from_messages_by_azure(self):
        """Get completion from messages by azure."""
        response = self.client.chat.completions.create(
            model=self.openai_api_model_azure,
            messages=self.messages,
            temperature=self.gpt_temperature, 
        )

        #print(str(response.choices[0].message))
        content = response.choices[0].message.content

        token_dict = {
        'prompt_tokens':response.usage.prompt_tokens,
        'completion_tokens':response.usage.completion_tokens,
        'total_tokens':response.usage.total_tokens,
        }
        return content, token_dict

    def _comletion_tokens(self):
        """Get comletion and tokens."""
        if self.azure:
            completion, token_dict = self._get_completion_from_messages_by_azure()
        else:
            completion, token_dict = self._get_completion_from_messages()
        self.translated_short_text = (
            completion
            .encode("utf8")
            .decode()
        )
        # Get the token usage from the API response
        self.total_tokens += token_dict['total_tokens']
        self.completion_tokens += token_dict['completion_tokens']
        self.prompt_tokens += token_dict['prompt_tokens']

    def _translate_text(self, content):
        """Translate the text."""
        # Call the OpenAI API for translation
        try:
            self.messages =  [
            {'role':'system', 
            'content': f"You are a translation assistant.Your task is to translate the content given to you by the user.{self.prompt}"},
            {'role': 'user',
            'content': f"{content}\n"},
            ]
            self._comletion_tokens()
        except Exception as e:
            # Time to wait for limitation of ChatGPT
            sleep_time = 60 * 3 + 5
            print(e, "\n"+f"Sleep {sleep_time} seconds.")
            time.sleep(sleep_time)
            self._comletion_tokens()

    def _translate_and_store(self, text):
        """Tranlate and store text."""
        if self.tlist:
            # Revert the replacement so that it can be judged whether the text has been translated
            text = self._text_replace_reverse(text)
            # If the text has been translated, return the translation result directly
            if text in self.translated_dict:
                self.translated_short_text = self.translated_dict[text]
            else:
                # Before translation, replace the text according to the transliteration table
                text = self._reverse_text_replace_reverse(text)
                # Else, call the translate_text function to translate and store the result in the dictionary
                self._translate_text(text)
                # Reverse the replacement of the transliteration table so than the text keeps the original content
                text = self._text_replace_reverse(text)
                self.translated_dict[text] = self.translated_short_text
                # Save the dictionary as a JSON file
                with open(self.jsonfile, "w", encoding="utf-8") as f:
                    json.dump(self.translated_dict, f, ensure_ascii=False, indent=4)
        else:
            # If the text has been translated, return the translation result directly
            if text in self.translated_dict:
                self.translated_short_text = self.translated_dict[text]
            else:
                # Else, call the translate_text function to translate and store the result in the dictionary
                self._translate_text(text)
                self.translated_dict[text] = self.translated_short_text
                # Save the dictionary as a JSON file
                with open(self.jsonfile, "w", encoding="utf-8") as f:
                    json.dump(self.translated_dict, f, ensure_ascii=False, indent=4)

    def _process_text(self):
        """Process the text."""
        # Replace all line breaks with spaces
        self.text = self.text.replace("\n", " ")
        # Replace multiple spaces with one space
        self.text = re.sub(r"\s+", " ", self.text)
        # If the transliteration table replacement is set, replace the text before translation
        if self.tlist:
            self._text_replace()
        # Split the text into short texts of no more than 1024 characters
        self._split_text()
        # If the test mode is turned on, only translate the first 3 short texts
        if self.test:
            self.short_text_list = self.short_text_list[:3]
        # Iterate through the short text list and translate each short text in turn
        for short_text in self.short_text_list:
            self.count += 1
            # Translate the current short text
            time.sleep(0.5)
            self._translate_and_store(short_text)
            short_text = self._replace_sign(short_text)
            self.translated_short_text = self._replace_sign(self.translated_short_text)
            short_text = self._text_replace_reverse(short_text)
            # Add the current short text and the translated text to the total text
            if self.bilingual_output.lower() == 'true':
                self.translated_text += f"{short_text}<br>\n{self.translated_short_text}<br>\n"
            else:
                self.translated_text += f"{self.translated_short_text}<br>\n"
            if self.show:
                print("*" * 3)
                print(short_text)
                print("*" * 1)
                print(self.translated_short_text)
                print("*" * 3)

    def _text_to_epub(self):
        """Write the translated text to the epub file."""
        text = self.translated_text.replace('\n', '<br>').replace("\n", "<br>")
        # Create an epub book object
        book = epub.EpubBook()
        # Set the metadata
        book.set_identifier(str(random.randint(100000, 999999)))
        book.set_title(self.title)
        book.set_language(self.language_code)
        # Create a chapter object
        c = epub.EpubHtml(title='Chapter 1', file_name='chap_1.xhtml', lang=self.language_code)
        c.content = text
        # Add the chapter to the book
        book.add_item(c)
        # Add the table of contents
        book.toc = (epub.Link('chap_1.xhtml', 'Chapter 1', 'chap_1'),)
        # Set spine order
        book.spine = ['nav', c]
        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        # Write the content to the epub book
        #print("\n" + text)
        try:
            epub.write_epub(self.new_filename, book, {})
        except Exception as e:
            print(f"Failed to write EPUB: {e}")

    def _get_title_of_md(self):
        """Get title of the md."""
        print("-" * 3)
        print("\033[1;32mINFO:Parsing the md title.\033[0m")
        with open(self.filename, 'r', encoding='utf-8') as file:
            for line in file:
                if line.startswith('#'):
                    self.title = line.replace('#', '').strip()
                    break
        print("-" * 3)
        print("\033[1;32mINFO:Finished parsing the md title.\033[0m")

    def _get_title_of_txt(self):
        """Get title of the txt."""
        print("-" * 3)
        print("\033[1;32mINFO:Parsing the txt title.\033[0m")
        title_extension = os.path.basename(self.filename)
        self.title = os.path.splitext(title_extension)[0]
        print("-" * 3)
        print("\033[1;32mINFO:Finished parsing the txt title.\033[0m")

    def _get_title_of_docx(self):
        """Get title of the docx."""
        try:
            print("-" * 3)
            print("\033[1;32mINFO:Parsing the docx file.\033[0m")
            with zipfile.ZipFile(self.filename) as zf:
                core_properties = etree.fromstring(zf.read("docProps/core.xml"))

            ns = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
                    "dc": "http://purl.org/dc/elements/1.1/",
                    "dcterms": "http://purl.org/dc/terms/",
                    "dcmitype": "http://purl.org/dc/dcmitype/",
                    "xsi": "http://www.w3.org/2001/XMLSchema-instance"}
            
            title_elements = core_properties.findall("dc:title", ns)
            if title_elements:
                self.title = title_elements[0].text
            else:
                self.title = "INFO:Unknown title."
            print("-" * 3)
            print("\033[1;32mINFO:Finished parsing the docx title.\033[0m")
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print("\033[91mERROR:Parsing the DOCX file.\033[0m")
            print("*" * 6)

    def _get_title_of_pdf(self):
        """Get title of the pdf."""
        try:
            print("-" * 3)
            print("\033[1;32mINFO:Parsing the pdf title.\033[0m") 
            with open(self.filename, 'rb') as file:
                parser = PDFParser(file)
                document = PDFDocument(parser)
                if 'Title' in document.info:
                    self.title = document.info['Title']
                else:
                    text = pdfminer.high_level.extract_text(file)
                    match = re.search(r'(?<=\n)([^\n]+)(?=\n)', text)
                    if match:
                        self.title = match.group(1)
                    else:
                        self.title = "INFO:Unknown title."
            print("-" * 3)
            print("\033[1;32mINFO:Finished parsing the pdf title.\033[0m") 
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print("\033[91mERROR:Parsing the pdf title.\033[0m")
            print("*" * 6)

    # step 1
    def get_title(self):
        """Get the title of file."""
        if self.filename.endswith('.pdf'):
            self._get_title_of_pdf()
            self._get_pdf_total_pages()
        elif self.filename.endswith('.txt'):
            self._get_title_of_txt()
        elif self.filename.endswith('.docx'):
            self._get_title_of_docx()
        elif self.filename.endswith('.mobi'):
            pass
        elif self.filename.endswith('.epub'):
            self.book = epub.read_epub(self.filename)
        elif self.filename.endswith('.md'):
            self._get_title_of_md()
        else:
            print("-" * 3)
            print("\033[91mINFO:Unsupported file type right now.\033[0m")
            print("-" * 3)
            sys.exit(0)

    def _get_md_content(self):
        """Get md content."""
        print("-" * 3)
        print("\033[1;32mINFO:Parsing the md content.\033[0m")
        with open(self.filename, 'r', encoding='utf-8') as file:
            self.text = file.read()

    def _get_txt_content(self):
        """Get txt content."""
        print("-" * 3)
        print("\033[1;32mINFO:Parsing the txt content.\033[0m")
        with open(self.filename, 'r', encoding='utf-8') as file:
            self.text = file.read()

    def _get_pdf_content(self):
        """Get pdf content."""
        try:
            print("-" * 3)
            print("\033[1;32mINFO:Parsing the pdf content.\033[0m")
            print("-" * 3)
            print(f"\033[1;32mINFO:Total pages of the pdf: {self.total_pages}\033[0m") 
            if self.end_page == -1:
                self.end_page = self.total_pages
            print("-" * 3)
            print(f"\033[1;32mINFO:Converting pdf from: Page {self.start_page} to Page {self.end_page}.\033[0m") 
            print("-" * 3)
            self._convert_pdf_to_text()
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print("\033[91mERROR:Parsing the pdf content.\033[0m")
            print("*" * 6)

    def _get_mobi_content(self):
        """Get mobi content."""
        try:
            print("-" * 3)
            print("\033[1;32mINFO:Parsing the mobi content.\033[0m")
            self._convert_mobi_to_text()
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print("\033[91mERROR:Parsing the MOBI content.\033[0m")
            print("*" * 6)

    def _get_epub_content(self):
        """Get mobi content."""
        try:
            print("-" * 3)
            print("\033[1;32mINFO:Parsing the EPUB content.\033[0m")
            self._convert_epub_to_text()
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print("\033[91mERROR:Parsing the EPUB content.\033[0m")
            print("*" * 6)

    # step 2
    def convert_text(self):
        """Convert the file ending with differnt types to text."""
        if self.filename.endswith('.pdf'):
            self._get_pdf_content()
        elif self.filename.endswith('.txt'):
            self._get_txt_content()
        elif self.filename.endswith('.mobi'):
            self._get_mobi_content()
        elif self.filename.endswith('.docx'):
            self._convert_docx_to_text()
        elif self.filename.endswith('.epub'):
            self._get_epub_content()
        elif self.filename.endswith('.md'):
            self._get_md_content()
        else:
            print("\033[91mINFO:Unsupported to access the content of this file type right now.\033[0m")

    # step 3
    def tranlate_file(self):
        """Translate the file."""
        if self.filename.endswith('.epub'):
            # Access all chapters of the epub file
            items = self.book.get_items()
            # Iterate through all chapters
            translated_all = ''
            print("-" * 3)
            print("\033[1;32mINFO:Translating the file content.\033[0m")
            for item in items:
                # If the chapter type is a document type, it needs to be translated
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Use BeautifulSoup to extract the original text
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    self.text = soup.get_text().strip()
                    img_html = ''
                    img_tags = soup.find_all('img')
                    for img_tag in img_tags:
                        img_html += str(img_tag) + '<br>'
                    # If the text is empty, skip this chapter
                    if not self.text:
                        continue
                    self._process_text()
                    # Replace the original chapter content with the translated text
                    item.set_content((img_html + self.translated_text.replace('\n', '<br>')).encode('utf-8'))
                    translated_all += self.translated_text
                    # If the test mode is turned on, only translate the first 3 chapters
                    if self.test and self.count >= 3:
                        break
            print("-" * 3)
            print("\033[1;32mINFO:Finished parsing and translating the file.\033[0m")
            # Write content to the epub file
            epub.write_epub(self.new_filename, self.book, {})
            # Write the translated text to the txt file
            with open(self.new_filenametxt, "w", encoding="utf-8") as f:
                f.write(translated_all.replace('<br>', ''))
        else:
            print("-" * 3)
            print("\033[1;32mINFO:Translating the file content.\033[0m")
            self._process_text()
            print("-" * 3)
            print("\033[1;32mINFO:Finished parsing and translating the file.\033[0m")
            print("-" * 3)
            # Write the translated text to the epub file
            print("\033[1;32mINFO:Writing the translated text to epub.\033[0m")  # 输出绿色的 "DEBUG"
            self._text_to_epub()
            # Write the translated text to the txt file
            print("-" * 3)
            print("\033[1;32mINFO:Writing the translated text to the txt file.\033[0m")
            with open(self.new_filenametxt, "w", encoding="utf-8") as f:
                f.write(self.translated_text.replace('<br>', ''))

    # step 4
    def caculate_tokens_costs(self):
        """Caculate the tokens."""
        cost = self.completion_tokens / 1000 * 0.002 + self.prompt_tokens / 1000 * 0.001
        print("-" * 3)
        print(f"\033[1;32mINFO:Use completion tokens: {self.completion_tokens}.\033[0m")
        print("-" * 3)
        print(f"\033[1;32mINFO:Use prompt tokens: {self.prompt_tokens}.\033[0m")
        print("-" * 3)
        print(f"\033[1;32mINFO:Use total tokens: {self.total_tokens}.\033[0m")
        print("-" * 3)
        print(f"\033[1;32mINFO:Total approximate cost: ${cost}.\033[0m")
        print("-" * 3)
        print(f"\033[1;34mINFO:Translation completed.\033[0m")
        print("-" * 3)

    # step 5
    def remove_jsonfile(self):
        """Remove the jsonfile."""
        try:
            os.remove(self.jsonfile)
            print(f"\033[1;34mFile '{self.jsonfile}' has been deleted.\033[0m")
            print("-" * 3)
        except Exception as e:
            print(f"An error occurred: {e}")
            print("*" * 6)
            print(f"\033[91mERROR:File '{self.jsonfile}' not found. No file was deleted.\033[0m")
            print("*" * 6)
